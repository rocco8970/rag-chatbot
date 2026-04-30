"""Document processing utilities for the RAG chatbot.

This module provides functions to extract and clean text from common
document formats used by the project: PDF, plain text, and DOCX.

All functions return a string. On error they return an empty string.
"""

from typing import List
import io
import re

import pdfplumber
import docx


def clean_extracted_text(text: str) -> str:
    """Clean extracted text.

    Cleaning steps:
    - Replace multiple spaces/tabs with a single space
    - Remove spaces that appear immediately before punctuation
    - Limit consecutive newlines to maximum two
    - Strip leading/trailing whitespace

    Args:
        text: Raw extracted text.

    Returns:
        Cleaned text. Returns empty string if input is falsy.
    """
    if not text:
        return ""

    try:
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Collapse multiple spaces/tabs into a single space
        text = re.sub(r"[ \t]+", " ", text)

        # Remove spaces before punctuation like .,;:!? )
        text = re.sub(r"\s+([\.,;:\!?\)])", r"\1", text)

        # Limit consecutive newlines to at most two
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip leading/trailing whitespace on the whole document
        text = text.strip()

        return text
    except Exception as e:
        print(f"clean_extracted_text error: {e}")
        return ""


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF bytes using pdfplumber.

    Handles multi-page PDFs and returns cleaned text. On error returns an
    empty string.

    Args:
        file_bytes: PDF file content as bytes.
        filename: Name of the file (used only for error messages).

    Returns:
        Extracted and cleaned text (str) or empty string on failure.
    """
    if not file_bytes:
        return ""

    try:
        text_parts: List[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_err:
                    # Continue processing other pages if a single page fails
                    print(f"Warning: failed to extract text from page {i} of {filename}: {page_err}")

        raw = "\n\n".join(text_parts)
        return clean_extracted_text(raw)
    except Exception as e:
        print(f"extract_text_from_pdf error ({filename}): {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode and return text content from a plain text file.

    Tries UTF-8 decoding first; on UnicodeDecodeError falls back to latin-1.
    Returns cleaned text. On error returns empty string.

    Args:
        file_bytes: File content as bytes.

    Returns:
        Cleaned text string or empty string on failure.
    """
    if not file_bytes:
        return ""

    try:
        try:
            decoded = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            # Fallback to latin-1 for legacy encodings
            decoded = file_bytes.decode("latin-1")

        return clean_extracted_text(decoded)
    except Exception as e:
        print(f"extract_text_from_txt error: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx.

    Extracts text from paragraphs and tables, preserving reasonable spacing.
    On error returns an empty string.

    Args:
        file_bytes: DOCX file content as bytes.

    Returns:
        Cleaned extracted text or empty string on failure.
    """
    if not file_bytes:
        return ""

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts: List[str] = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        # Tables: iterate rows/cells and collect text
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        raw = "\n\n".join(parts)
        return clean_extracted_text(raw)
    except Exception as e:
        print(f"extract_text_from_docx error: {e}")
        return ""


__all__ = [
    "clean_extracted_text",
    "extract_text_from_pdf",
    "extract_text_from_txt",
    "extract_text_from_docx",
]


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200):
    """Split `text` into overlapping chunks.

    Logic:
    - Splits text into chunks of approximately `chunk_size` characters.
    - Prefers breaking at sentence boundaries ('. ' or '\n') when a boundary
      is found after 50% of `chunk_size` and before the chunk end.
    - Ensures each chunk overlaps the previous by `overlap` characters.

    Args:
        text: Source text to chunk.
        chunk_size: Target chunk size in characters (must be > 0).
        overlap: Number of characters to overlap between consecutive chunks.

    Returns:
        A list of dicts with keys: `content` (str), `start_pos` (int), `end_pos` (int).
        Returns an empty list for empty input.
    """
    from typing import List, Dict

    results: List[Dict[str, object]] = []
    if not text:
        return results

    if chunk_size <= 0:
        raise ValueError("chunk_size must be > 0")
    if overlap < 0:
        raise ValueError("overlap must be >= 0")

    n = len(text)
    # Ensure overlap is smaller than chunk_size
    if overlap >= chunk_size:
        overlap = max(0, chunk_size // 10)

    start = 0
    while start < n:
        # end of the candidate window
        window_end = min(start + chunk_size, n)

        # If this is the last chunk, take the remainder
        if window_end == n:
            end_pos = n
        else:
            # Try to find a sentence boundary between halfway and window_end
            min_break = start + chunk_size // 2
            break_pos = None
            # Prefer newline then sentence-ending with space
            for sep in ['\n', '. ']:
                idx = text.rfind(sep, min_break, window_end)
                if idx != -1:
                    candidate = idx + len(sep)
                    if candidate <= window_end:
                        if break_pos is None or candidate > break_pos:
                            break_pos = candidate

            if break_pos is not None and break_pos > start:
                end_pos = break_pos
            else:
                # no good boundary found; just use window_end
                end_pos = window_end

        # Extract and clean chunk content but preserve positions based on original text
        chunk_raw = text[start:end_pos]
        chunk_content = clean_extracted_text(chunk_raw)

        results.append({
            "content": chunk_content,
            "start_pos": int(start),
            "end_pos": int(end_pos),
        })

        if end_pos >= n:
            break

        # Advance start by chunk end minus overlap, ensure progress
        next_start = end_pos - overlap
        if next_start <= start:
            # fallback: advance by chunk_size - overlap (at least 1)
            next_start = start + max(1, chunk_size - overlap)

        start = next_start

    return results


if __name__ == '__main__':
    # Quick test examples for chunk_text
    small = "This is a short document. It has two sentences."  # < chunk_size
    print('\nExample 1: short text')
    print(chunk_text(small, chunk_size=100, overlap=20))

    long_sentences = (
        "Sentence one. " * 10 + "\n" + "Sentence two is a bit longer. " * 8 + "\n" + "Final short sentence."
    )
    print('\nExample 2: longer text')
    chunks = chunk_text(long_sentences, chunk_size=200, overlap=50)
    print(f"Produced {len(chunks)} chunks")
    for i, c in enumerate(chunks[:3]):
        print(f"--- chunk {i} ({c['start_pos']}:{c['end_pos']}) ---")
        print(c['content'][:200])

